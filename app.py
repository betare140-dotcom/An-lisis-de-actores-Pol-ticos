import io
import json
import os
import re
import unicodedata
import zipfile
from datetime import datetime
import google.generativeai as genai
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
import pandas as pd
import streamlit as st

MESES_ES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril", 5: "mayo", 6: "junio",
    7: "julio", 8: "agosto", 9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
}

st.set_page_config(
    page_title="Monitoreo Político Oficial",
    page_icon="📊",
    layout="centered"
)

st.title("📊 Generador de Monitoreo de Actores Políticos")
st.write(
    "Sistema universal de monitoreo político para cualquier perfil o cargo público. "
    "Permite procesar archivos individuales o fusionar múltiples archivos (Radio/TV, Portales Web y Redes) en un solo reporte oficial consolidado."
)

# API Key Pre-integrada
GEMINI_API_KEY = "AQ.Ab8RN6LoOHgBblHSIETp2LjyBofO48YsSqSeojXYFAAKGvFa0w"
genai.configure(api_key=GEMINI_API_KEY)

# ==============================================================================
# BASE DE CONOCIMIENTO Y CRITERIOS UNIVERSALES DE CLASIFICACIÓN POLÍTICA
# ==============================================================================
SYSTEM_PROMPT_UNIVERSAL = """
Eres un analista senior de inteligencia política, comunicación gubernamental y control de crisis.
Tu objetivo es clasificar publicaciones para CUALQUIER actor político objetivo evaluando el texto desde la perspectiva DIRECTA del actor político y su equipo de control de daños.

Determina el impacto reputacional del texto para el actor político en una de dos categorías: 'POSITIVA' o 'NEGATIVA'.

==============================================================================
REGLAS Y CRITERIOS DE CLASIFICACIÓN (PERSPECTIVA DEL ACTOR POLÍTICO):
==============================================================================

1. CLASIFICA ESTRICTAMENTE COMO 'NEGATIVA' (Cualquier crisis, crítica, queja o afectación):
   - ESCÁNDALOS POLICIALES Y ABUSO DE FUERZA: Elementos de policía o tránsito disparando, encañonando, extorsionando, agrediendo automovilistas, cobro de cuotas, prepotencia, o calificados como 'asesinos', 'delincuentes' o 'as3sinos'.
   - INVESTIGACIONES, AUDITORÍAS Y ÓRDENES JUDICIALES: Órdenes de aprehensión, denuncias penales, observaciones de la Auditoría Superior (ASE/ASF), dinero o recursos no comprobados/no justificados, daño patrimonial o desvíos contra el actor, su pareja o funcionarios.
   - CRÍTICA POLÍTICA, VIDEOS Y COLUMNAS DE OPINIÓN: Expresiones editoriales de advertencia o golpeteo político como "se viene la noche", "bajo la lupa", "pone en jaque", "en el ojo del huracán", "focos rojos", "alerta", "sinvergüenza", "fichita" o acusaciones de traición.
   - ACUSACIONES DE OPACIDAD Y FALTA DE TRANSPARENCIA: Críticas de adversarios, regidores u organizaciones que señalen falta de claridad en el manejo de recursos, licitaciones o contratos.
   - DELINCUENCIA E INSEGURIDAD EN SU TERRITORIO: Notas sobre robos (bicicletas, casas, comercios, autopartes), asaltos al transporte, homicidios o balaceras en su municipio, especialmente si se destaca la falta de vigilancia o cercanía a oficinas gubernamentales.
   - RECLAMOS CIUDADANOS Y SERVICIOS DEFICIENTES: Quejas por baches, falta de agua, luminarias descompuestas, basura o abandono en colonias y juntas auxiliares.
   - HASHTAGS DE ATAQUE O LENGUAJE EVASIVO: Publicaciones con etiquetas como #peligro, #corrupcion, #balazos, #as3sinos, #orden, #aprehension, #escandalo, #rateros.

2. CLASIFICA COMO 'POSITIVA' / INFORMATIVA (Beneficio, Cobertura o Agenda):
   - OBRAS PÚBLICAS Y EVENTOS INSTITUCIONALES: Arranques de obra, pavimentación, alumbrado, cursos de verano, eventos culturales, ventanilla digital, apoyos sociales, despensas y ferias del empleo.
   - DECLARACIONES Y POSTURA OFICIAL: Difusión de sus discursos, entrevistas, réplicas, comunicados de prensa o iniciativas legislativas.
   - AGENDA PÚBLICA Y SERVICIOS: Convocatorias, actividades cívicas, trámites y servicios institucionales.
   - RESPALDOS POLÍTICOS Y ENCUESTAS: Felicitaciones, cierres de filas partidistas o posicionamientos favorables en sondeos.
"""

model_clasificador = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    system_instruction=SYSTEM_PROMPT_UNIVERSAL,
    generation_config={"temperature": 0.0}
)

model_redactor = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    generation_config={"temperature": 0.2}
)

def quitar_acentos(texto):
    if texto is None or pd.isna(texto):
        return ""
    return ''.join(c for c in unicodedata.normalize('NFD', str(texto)) if unicodedata.category(c) != 'Mn').lower()

def normalizar_cadena(texto):
    t = quitar_acentos(texto)
    return re.sub(r'[^a-z0-9]', '', t)

def normalizar_nombre_candidato(nombre):
    t = quitar_acentos(str(nombre).strip())
    t_clean = re.sub(r'[^a-z0-9]', '', t)
    
    alias_map = {
        "pepechedraui": "JOSÉ CHEDRAUI BUDIB",
        "josechedraui": "JOSÉ CHEDRAUI BUDIB",
        "rmvb": "RAFAEL MORENO VALLE BUITRÓN",
        "rafamorenovalle": "RAFAEL MORENO VALLE BUITRÓN",
        "rafaelmorenovalle": "RAFAEL MORENO VALLE BUITRÓN",
        "carolinabeau": "CAROLINA BEAUREGARD MARTÍNEZ",
        "carolinabeauregard": "CAROLINA BEAUREGARD MARTÍNEZ",
        "genovevahuerta": "GENOVEVA HUERTA VILLEGAS",
        "gabrielasanchez": "GABRIELA SÁNCHEZ SAAVEDRA",
        "lauraartemisa": "LAURA ARTEMISA GARCÍA CHÁVEZ",
        "lupitacuautle": "GUADALUPE CUAUTLE TORRES",
        "guadalupecuautle": "GUADALUPE CUAUTLE TORRES",
        "tonantzinfernandez": "TONANTZIN FERNÁNDEZ DÍAZ",
        "lizsanchez": "LIZETH SÁNCHEZ GARCÍA",
        "lizethsanchez": "LIZETH SÁNCHEZ GARCÍA",
        "nestorcamarillo": "NESTOR CAMARILLO MEDINA",
        "celinapena": "CELINA PEÑA GUZMÁN",
        "rodrigoabdala": "RODRIGO ABDALA DARTIGUES",
        "delfinapozos": "DELFINA POZOS VERGARA",
        "xitlalicceja": "XITLALIC CEJA GARCÍA",
        "blancaalcala": "BLANCA ALCALÁ RUIZ"
    }
    for alias, canon in alias_map.items():
        if alias in t_clean or t_clean in alias:
            return canon
    return str(nombre).strip().upper()

# --- PATRONES INSTITUCIONALES EXPANDIDOS ---
PATRONES_INSTITUCIONALES_GENERICOS = [
    r'ayuntamiento', r'gobierno', r'gobpue', r'gobiernoded', r'ayto', r'snandresoficial', r'cholulaoficial',
    r'secretaria', r'dependencia', r'organismo', r'sindicatura', r'presidencia', r'comunicacionsocial',
    r'seguridadciudadana', r'seguridadpublica', r'proteccioncivil', r'policiamunicipal', r'policiastat',
    r'serviciospublicos', r'servicios_pub', r'obraspublicas', r'desarrollourbano', r'desarrolloeconomico',
    r'dif', r'sistemadif', r'difestatal', r'difmunicipal', r'institutodelamujer', r'institutodelajuventud',
    r'organismodeagua', r'serviciodelimpia', r'organismolimpia', r'derechoshumanos', r'casadecultura',
    r'bienestarsanandres', r'bienestarcholula', r'ssppc', r'ssppcsnandres'
]

ABREV_APELLIDOS = {
    'fdz': 'fernandez', 'hdez': 'hernandez', 'mtz': 'martinez', 'glez': 'gonzalez',
    'gcia': 'garcia', 'lpz': 'lopez', 'prz': 'perez', 'sdo': 'saavedra', 'chvz': 'chavez'
}

def es_cuenta_del_actor_universal(autor, handle, actor_target):
    aut = normalizar_cadena(autor)
    hnd = normalizar_cadena(handle)
    act = normalizar_cadena(actor_target)
    
    if not act:
        return False
        
    if act in aut or act in hnd or aut in act:
        return True
        
    tokens_actor = [t for t in re.findall(r'\w+', quitar_acentos(actor_target)) if len(t) > 2]
    for abrev, apellido_completo in ABREV_APELLIDOS.items():
        if apellido_completo in tokens_actor:
            tokens_actor.append(abrev)
            
    nombres_distintivos = [t for t in tokens_actor if len(t) >= 5 and t not in ['perez', 'lopez', 'garcia', 'martinez']]
    for n in nombres_distintivos:
        if n in aut or n in hnd:
            otros_tokens = [t for t in tokens_actor if t != n]
            if any(ot in aut or ot in hnd for ot in otros_tokens):
                return True
            if hnd == n or aut == n:
                return True
                
    coincidencias_aut = sum(1 for t in tokens_actor if t in aut)
    coincidencias_hnd = sum(1 for t in tokens_actor if t in hnd)
    if coincidencias_aut >= 2 or coincidencias_hnd >= 2:
        return True
        
    return False

def es_institucional_universal(autor, handle, texto_post=""):
    texto = (str(autor) + " " + str(handle))
    texto_norm = normalizar_cadena(texto)
    for p in PATRONES_INSTITUCIONALES_GENERICOS:
        if re.search(p, texto_norm):
            return True
            
    t_clean = quitar_acentos(texto_post).lower()
    frases_institucionales = [
        'nuestra presidenta municipal', 'nuestra presidenta honoraria', 'nuestro presidente municipal',
        'mi mas sincero agradecimiento', 'nuestro compromiso de seguir', 'desde el gobierno municipal',
        'agradecemos a nuestra presidenta', 'agradecemos a nuestro presidente', 'los invitamos a participar en nuestro'
    ]
    if any(f in t_clean for f in frases_institucionales):
        return True
        
    return False

def limpiar_dataframe_redes_automatico(df_raw, actor_nombre_target):
    def es_descartable(row):
        autor = obtener_campo(row, ["Autor", "Author name", "Author", "User", "Username", "Handle", "Fuente", "Nombre del Medio", "Canal"])
        handle = obtener_campo(row, ["Author handle (@username)", "Handle", "Username", "Screen Name", "Account", "Perfil"])
        detalle = obtener_campo(row, ["Contenido", "Detail", "Titulo", "Título", "Summary", "Síntesis", "Nota"])
        link = obtener_campo(row, ["Link URL Medio", "URL", "Enlace", "Link", "Link de Nota"])
        
        # 1. Comprobar cuenta del actor político
        if es_cuenta_del_actor_universal(autor, handle, actor_nombre_target):
            return True
            
        # 2. Comprobar cuenta institucional / DIF / Ayuntamiento
        if es_institucional_universal(autor, handle, detalle):
            return True
            
        # 3. Comprobar si el link del post apunta al perfil propio del candidato
        if link and actor_nombre_target:
            link_norm = normalizar_cadena(link)
            tokens_actor = [t for t in re.findall(r'\w+', quitar_acentos(actor_nombre_target)) if len(t) >= 5]
            for n in tokens_actor:
                if f"x.com/{n}" in link.lower() or f"twitter.com/{n}" in link.lower() or f"facebook.com/{n}" in link.lower() or f"{n}fdz" in link_norm or f"{n}t" in link_norm:
                    return True
        
        det_lower = quitar_acentos(detalle)
        if any(p in det_lower for p in ['mis ahijados', 'con toda la actitud #graciasdios', 'primeracomunion', 'en familia festejando']):
            return True
            
        return False

    mask_descarte = df_raw.apply(es_descartable, axis=1)
    df_limpio = df_raw[~mask_descarte].copy()
    total_descartadas = mask_descarte.sum()
    
    return df_limpio, total_descartadas

def cargar_archivo_seguro(file):
    try:
        file_bytes = file.read()
        file.seek(0)
    except Exception:
        file_bytes = file

    # 1. Intentar como Excel usando openpyxl
    try:
        return pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
    except Exception:
        pass

    # 2. Intentar como CSV UTF-8
    try:
        df_csv = pd.read_csv(io.BytesIO(file_bytes), on_bad_lines="skip")
        return {"Reporte": df_csv}
    except Exception:
        pass

    # 3. Intentar como CSV Latin-1
    try:
        df_csv = pd.read_csv(io.BytesIO(file_bytes), encoding="latin1", on_bad_lines="skip")
        return {"Reporte": df_csv}
    except Exception:
        pass

    return {}

def obtener_campo(row, lista_cols):
    for c in lista_cols:
        if c in row.index:
            val = row[c]
            if val is not None and not pd.isna(val):
                v = str(val).strip()
                if v and v.lower() not in ["nan", "none", "null"]:
                    return v
        for col_existente in row.index:
            col_str = str(col_existente).strip()
            if col_str.lower() == c.lower() or quitar_acentos(col_str) == quitar_acentos(c):
                val = row[col_existente]
                if val is not None and not pd.isna(val):
                    v = str(val).strip()
                    if v and v.lower() not in ["nan", "none", "null"]:
                        return v
    return ""

def obtener_columna_serie(df_data, lista_posibles_cols):
    for c in lista_posibles_cols:
        if c in df_data.columns:
            return df_data[c]
        for col_existente in df_data.columns:
            col_str = str(col_existente).strip()
            if col_str.lower() == c.lower() or quitar_acentos(col_str) == quitar_acentos(c):
                return df_data[col_existente]
    return pd.Series([""] * len(df_data), index=df_data.index)

def parsear_fecha_perfecta(val):
    if not val or pd.isna(val) or str(val) == "nan":
        return pd.NaT

    s = str(val).strip()
    if "," in s:
        s = s.split(",")[0].strip()
    s_date = s.split(" ")[0].strip()

    if "-" in s_date:
        try:
            parts = s_date.split("-")
            if len(parts) == 3 and len(parts[0]) == 4:
                return datetime(int(parts[0]), int(parts), int(parts))
            elif len(parts) == 3 and len(parts) == 4:
                return datetime(int(parts), int(parts), int(parts[0]))
        except Exception:
            pass

    if "/" in s_date:
        try:
            parts = s_date.split("/")
            if len(parts) == 3:
                if len(parts[0]) == 4:
                    return datetime(int(parts[0]), int(parts), int(parts))
                else:
                    d = int(parts[0])
                    m = int(parts)
                    y = int(parts)
                    if y < 100: y += 2000
                    return datetime(y, m, d)
        except Exception:
            pass

    try:
        return pd.to_datetime(s_date, dayfirst=True, errors="coerce")
    except Exception:
        return pd.NaT

def limpiar_texto(texto):
    if not isinstance(texto, str) or texto == "nan":
        return ""
    texto_sin_urls = re.sub(r"https?://\S+", "", texto)
    texto_sin_emojis = re.sub(r":[a-zA-Z0-9_\-|]+:", "", texto_sin_urls)
    texto_sin_emojis = re.sub(r'[\U00010000-\U0010ffff]', '', texto_sin_emojis)
    lineas = [
        re.sub(r"[ \t]+", " ", line).strip()
        for line in texto_sin_emojis.split("\n")
        if line.strip()
    ]
    return "\n".join(lineas)

def estandarizar_categoria_medio(m_type_raw):
    if m_type_raw is None or pd.isna(m_type_raw):
        return "PORTALES DIGITALES"
    t_norm = quitar_acentos(str(m_type_raw).strip())
    if "tele" in t_norm or "tv" in t_norm:
        return "TELEVISIÓN"
    elif "rad" in t_norm or "fm" in t_norm or "am" in t_norm:
        return "RADIO"
    elif any(k in t_norm for k in ["portal", "web", "online", "internet", "digital", "comun"]):
        return "PORTALES DIGITALES"
    elif any(k in t_norm for k in ["prensa", "periodico", "diario", "impreso"]):
        return "PRENSA LOCAL"
    elif any(k in t_norm for k in ["columna", "opinion"]):
        return "COLUMNAS"
    elif any(k in t_norm for k in ["redes", "social", "twitter", "facebook", "instagram", "tiktok", "youtube", "x"]):
        return "REDES SOCIALES"
    return "PORTALES DIGITALES"

def reparar_desfase_columnas_excel(df):
    if df is None or df.empty:
        return df
        
    cols = [str(c).strip() for c in df.columns]
    
    es_desfasado = False
    if "Hora" in cols:
        sample_hora = df["Hora"].dropna().astype(str).head(10).tolist()
        if any(re.match(r'^(Puebla|M[eé]xico|Tlaxcala|Veracruz|CDMX|Hidalgo|Nacional|Internacional)$', v.strip(), re.I) for v in sample_hora):
            es_desfasado = True
            
    if "Alcance" in cols and not es_desfasado:
        sample_alcance = df["Alcance"].dropna().astype(str).head(10).tolist()
        if any(v.startswith("http") for v in sample_alcance):
            es_desfasado = True

    if es_desfasado:
        columnas_reales_ordenadas = [
            "ID Nota", "Menu", "Titulo", "Autor", "Fecha", 
            "Estado", "Pais", "Nombre del Medio", "Tipo de Medio", 
            "Tipo de Nota", "Sentimiento", "Costo", "Alcance", 
            "Link URL Medio", "Link de Nota"
        ]
        
        num_cols = len(df.columns)
        df_reparado = pd.DataFrame()
        
        for idx, col_name in enumerate(columnas_reales_ordenadas):
            if idx < num_cols:
                df_reparado[col_name] = df.iloc[:, idx].values
            else:
                df_reparado[col_name] = ""
                
        df_reparado["Hora"] = ""
        return df_reparado
        
    return df

def clasificar_lote_con_ia(lista_notas, actor_nombre):
    prompt = f"""
Clasifica las siguientes publicaciones respecto al actor político: "{actor_nombre}".
Aplica estrictamente los criterios de evaluación desde la perspectiva directa del actor político y control de daños.

NOTAS A EVALUAR:
{json.dumps(lista_notas, ensure_ascii=False)}

Responde ÚNICAMENTE un JSON válido con este formato exacto:
[
  {{"id": 0, "sentimiento": "POSITIVA"}},
  {{"id": 1, "sentimiento": "NEGATIVA"}}
]
"""
    try:
        response = model_clasificador.generate_content(prompt)
        raw_txt = response.text.strip()
        raw_txt = re.sub(r"^```json\s*", "", raw_txt, flags=re.I)
        raw_txt = re.sub(r"^```\s*", "", raw_txt)
        raw_txt = re.sub(r"\s*```$", "", raw_txt)
        
        datos = json.loads(raw_txt)
        res_map = {}
        for item in datos:
            sent = item.get("sentimiento", "POSITIVA").upper()
            res_map[item["id"]] = "NEGATIVA" if "NEGAT" in sent else "POSITIVA"
        return res_map
    except Exception:
        res_map = {}
        patrones_negativos_regex = [
            r'as3sin[oa]s?', r'asesin[oa]s?', r'orden(es)? de aprehensi[oó]n', r'no justificad[oa]s?', r'se viene la noche',
            r'extorsi[oó]n', r'extorsiona', r'disparan?do?', r'encañonan?', r'balaz[o0]s?', r'arma de fuego',
            r'cortando cartucho', r'vidrio roto', r'agresi[oó]n armada', r'polic[ií]as corrupt[oa]s?', r'prepoten(te|cia)',
            r'inseguridad golpea', r'roban? ', r'robo a ', r'bajo la lupa', r'opacidad', r'en jaque', r'esc[aá]ndalo',
            r'daño patrimonial', r'desv[ií]o', r'auditor[ií]a', r'nepotismo', r'desfalco', r'sinverg[uü]enza', r'fichita',
            r'baches?', r'falta de agua', r'sin agua', r'luz apagada', r'luminarias? descompuesta', r'basura acumulada'
        ]
        for item in lista_notas:
            t = str(item.get("texto", "")).lower()
            es_neg = any(re.search(p, t) for p in patrones_negativos_regex)
            res_map[item["id"]] = "NEGATIVA" if es_neg else "POSITIVA"
        return res_map

def determinar_sentimiento_df(df_data, actor_nombre_target, es_tradicionales):
    if es_tradicionales:
        sent_col_name = None
        for col_name in df_data.columns:
            if quitar_acentos(str(col_name)) in ["sentimiento", "sentiment", "sentimiento de la nota", "tono", "sentimiento nota"]:
                sent_col_name = col_name
                break
        if sent_col_name:
            sent_series = df_data[sent_col_name].fillna("").astype(str).str.lower()
            sentimientos_lista = []
            patrones_crisis_forzada = [
                r'as3sin[oa]s?', r'orden(es)? de aprehensi[oó]n', r'no justificad[oa]s?', r'se viene la noche',
                r'extorsi[oó]n', r'disparan?do?', r'encañonan?', r'agresi[oó]n armada', r'inseguridad golpea'
            ]
            for _, row in df_data.iterrows():
                texto_row = str(obtener_campo(row, ["Contenido", "Detail", "Titulo", "Título", "Summary", "Síntesis", "Nota"])).lower()
                sent_raw = str(row[sent_col_name]).lower() if sent_col_name in row.index else ""
                if any(re.search(p, texto_row) for p in patrones_crisis_forzada) or any(k in sent_raw for k in ["negat", "critica", "contra"]):
                    sentimientos_lista.append("NEGATIVA")
                else:
                    sentimientos_lista.append("POSITIVA")
            return sentimientos_lista

    # Evaluación de Redes Sociales con perspectiva de IA del actor político
    df_eval = df_data.reset_index(drop=True)
    sentimientos_finales = ["POSITIVA"] * len(df_eval)
    lote_tamano = 15
    total_lotes = (len(df_eval) + lote_tamano - 1) // lote_tamano
    
    progreso = st.progress(0)

    for l_idx in range(total_lotes):
        sub_df = df_eval.iloc[l_idx * lote_tamano : (l_idx + 1) * lote_tamano]
        lista_lote = []
        for local_id, (_, row) in enumerate(sub_df.iterrows()):
            texto = obtener_campo(row, ["Contenido", "Detail", "Titulo", "Título", "Summary", "Síntesis", "Sintesis", "Title", "Encabezado", "Tema", "Nota"])
            lista_lote.append({"id": local_id, "texto": texto})
            
        res_map = clasificar_lote_con_ia(lista_lote, actor_nombre_target)
        for local_id in range(len(sub_df)):
            real_idx = l_idx * lote_tamano + local_id
            if real_idx < len(sentimientos_finales):
                sentimientos_finales[real_idx] = res_map.get(local_id, "POSITIVA")
            
        progreso.progress((l_idx + 1) / total_lotes)
        
    progreso.empty()
    return sentimientos_finales

# --- GENERADOR UNIVERSAL DEL RESUMEN EJECUTIVO EN PÁRRAFOS CONCISOS ---
def limpiar_texto_para_resumen(texto):
    if not isinstance(texto, str) or texto == "nan":
        return ""
    t = re.sub(r'https?://\S+', '', texto)
    t = re.sub(r'---\s*transcripci[oó]n\s*---[\s\S]*', '', t, flags=re.I)
    t = re.sub(r'kind:\s*captions.*', '', t, flags=re.I)
    t = re.sub(r'#([a-zA-Z0-9_]+)', r'\1', t)
    t = re.sub(r'[\U00010000-\U0010ffff]', '', t)
    t = re.sub(r':[a-zA-Z0-9_\-|]+:', '', t)
    t = re.sub(r'\s+', ' ', t).strip()
    return t

def extraer_resumen_temas_real(df_data, actor_nombre):
    pos_df = df_data[df_data["sentimiento_final"].isin(["POSITIVA", "NEUTRA"])]
    neg_df = df_data[df_data["sentimiento_final"] == "NEGATIVA"]

    pos_textos = []
    columnas_posibles_texto = ["Titulo", "Título", "Contenido", "Detail", "Summary", "Síntesis", "Sintesis", "Encabezado", "Tema", "Nota", "Title"]
    
    for col in columnas_posibles_texto:
        if col in pos_df.columns or any(quitar_acentos(str(c)) == quitar_acentos(col) for c in pos_df.columns):
            serie_t = obtener_columna_serie(pos_df, [col])
            raw_list = serie_t.dropna().astype(str).tolist()
            pos_textos = [limpiar_texto_para_resumen(t) for t in raw_list if len(limpiar_texto_para_resumen(t)) > 10 and "teaser" not in t.lower()]
            if len(pos_textos) > 0:
                break

    neg_textos = []
    for col in columnas_posibles_texto:
        if col in neg_df.columns or any(quitar_acentos(str(c)) == quitar_acentos(col) for c in neg_df.columns):
            serie_t = obtener_columna_serie(neg_df, [col])
            raw_list = serie_t.dropna().astype(str).tolist()
            neg_textos = [limpiar_texto_para_resumen(t) for t in raw_list if len(limpiar_texto_para_resumen(t)) > 10 and "teaser" not in t.lower()]
            if len(neg_textos) > 0:
                break

    # Deduplicar textos base
    pos_unicos = list(dict.fromkeys(pos_textos))
    neg_unicos = list(dict.fromkeys(neg_textos))

    prompt = f"""
Eres un analista senior de comunicación política y redacción ejecutiva.
Redacta el "RESUMEN" ejecutivo oficial para el actor político: "{actor_nombre}".

REGLAS DE FORMATO Y ESTILO (OBLIGATORIAS):
1. CERO DUPLICADOS: Agrupa las notas por tema. Si varias notas tratan del mismo hecho (ej. un incidente vial con policías, una columna de opinión, un robo o una obra pública), redacta UN SOLO punto que sintetice el caso globalmente.
2. REDACCIÓN DESCRIPTIVA EJECUTIVA: Cada punto debe ser un párrafo fluido, profesional y descriptivo (de 2 a 3 líneas), explicando los hechos con nombres de programas, lugares, vialidades o instituciones.
3. PROHIBIDO COPIAR Y PEGAR TITULARES O FRAGMENTOS LITERALES: Debes sintetizar y redactar con tus propias palabras en tono institucional y formal (igual que en los temas positivos).
4. ESTRUCTURA EXACTA DE CADA PUNTO:
   [Número]. [Título del Eje Temático en Mayúsculas y Minúsculas]: [Descripción ejecutiva redactada formalmente].
5. CANTIDAD DE PUNTOS:
   - Genera de 1 a 3 ejes temáticos en 'Temas relevantes informativos'.
   - Genera de 1 a 3 ejes temáticos en 'Temas negativos'.
   - Si no hay notas negativas ({len(neg_unicos)} negativas registradas), escribe obligatoriamente:
     1. No se registraron temas negativos en el periodo analizado.

EJEMPLO DE ESTRUCTURA:
Temas relevantes informativos
1. Promoción del Desarrollo Económico: La administración presentó la Ventanilla Digital de Inversiones ante el Consejo Directivo de CANACINTRA Puebla para agilizar trámites y atraer nuevas empresas.
2. Actividades Culturales y Recreativas: Se llevó a cabo la clausura del Curso de Verano 2026 en la Casa de Cultura Tlanezcalli, enfocado en el desarrollo artístico de niñas, niños y jóvenes.

Temas negativos
1. Incidentes de Seguridad Vial: Cobertura mediática y reclamos ciudadanos tras el presunto intento de extorsión y agresión atribuido a policías municipales contra una pareja en el Bulevar del Niño Poblano.
2. Controversias y Crítica Política: Publicación de videocolumnas de opinión y señalamientos sobre investigaciones administrativas y observaciones de recursos públicos.

NOTAS POSITIVAS DISPONIBLES:
{json.dumps(pos_unicos[:25], ensure_ascii=False)}

NOTAS NEGATIVAS DISPONIBLES:
{json.dumps(neg_unicos[:25], ensure_ascii=False)}
"""
    try:
        response = model_redactor.generate_content(prompt)
        res_ia = response.text.strip()
        if len(neg_unicos) > 0 and "no se registraron temas negativos" in res_ia.lower():
            pass
        elif res_ia and len(res_ia) > 30 and len(res_ia) < 1800:
            if not res_ia.startswith("Temas relevantes informativos"):
                res_ia = "Temas relevantes informativos\n" + res_ia
            return res_ia
    except Exception:
        pass

    # Fallback inteligente y descriptivo sin duplicados
    lineas_res = ["Temas relevantes informativos"]
    if len(pos_unicos) > 0:
        for i, t in enumerate(pos_unicos[:3], 1):
            t_clean = t.split('.')[0].strip()
            lineas_res.append(f"{i}. Gestión y Agenda Institucional: Seguimiento y difusión a {t_clean.lower()}.")
    else:
        lineas_res.append("1. Agenda Institucional: Difusión de actividades públicas y agenda de trabajo.")

    lineas_res.append("\nTemas negativos")
    if len(neg_unicos) > 0:
        ejes_neg_detectados = []
        t_todo = " ".join(neg_unicos).lower()
        if any(k in t_todo for k in ["policía", "policia", "tránsito", "disparando", "extorsión", "niño poblano", "as3sinos"]):
            ejes_neg_detectados.append("Incidentes de Seguridad y Actuación Policial: Cobertura mediática y quejas ciudadanas sobre presuntos abusos, agresiones y revisiones a automovilistas por parte de elementos de seguridad vial.")
        if any(k in t_todo for k in ["columna", "contrastes", "aprehensión", "no justificados", "se viene la noche", "opacidad", "transparencia"]):
            ejes_neg_detectados.append("Crítica Política y Fiscalización: Publicación de columnas editoriales y señalamientos respecto a investigaciones administrativas, transparencia y manejo de recursos públicos.")
        if any(k in t_todo for k in ["robo", "asalt", "bicicleta", "inseguridad golpea"]):
            ejes_neg_detectados.append("Demandas Ciudadanas de Seguridad: Denuncias vecinales sobre hechos delictivos y robo en zonas céntricas, exigiendo mayor patrullaje y prevención del delito.")
            
        if not ejes_neg_detectados:
            for t in neg_unicos[:2]:
                t_clean = t.split('.')[0].strip()
                ejes_neg_detectados.append(f"Controversias y Señalamientos Públicos: Cobertura sobre {t_clean.lower()}.")
                
        for i, eje_txt in enumerate(ejes_neg_detectados[:3], 1):
            lineas_res.append(f"{i}. {eje_txt}")
    else:
        lineas_res.append("1. No se registraron temas negativos en el periodo analizado.")

    return "\n".join(lineas_res)

def obtener_link_inteligente(row):
    urls_en_fila = []
    
    # 1. Buscar en campos específicos de enlace
    campos_prioridad = ["Link URL Medio", "Link de Nota", "URL", "Enlace", "Link", "Link Medio", "Alcance"]
    for c in campos_prioridad:
        val = obtener_campo(row, [c])
        if val and str(val).startswith("http") and str(val).strip() not in urls_en_fila:
            urls_en_fila.append(str(val).strip())
            
    # 2. Buscar en todas las celdas de la fila por si hubo desplazamiento humano en Excel
    for val in row.values:
        if val is not None and not pd.isna(val):
            s_val = str(val).strip()
            if s_val.startswith("http") and s_val not in urls_en_fila:
                urls_en_fila.append(s_val)
                
    if not urls_en_fila:
        return ""
        
    # Filtrar descartando siempre enlaces de Testigo
    urls_sin_testigo = [u for u in urls_en_fila if "hanakua.mx/Testigo" not in u]
    
    # Regla A: Si hay una URL externa del medio (Portal Web, Prensa o Red Social), tiene prioridad máxima
    urls_medios_externos = [u for u in urls_sin_testigo if "hanakua.mx" not in u.lower()]
    if urls_medios_externos:
        return urls_medios_externos[0]
        
    # Regla B: Si es Radio / TV (no hay web externa), usar el link de la nota en Hanakua ([https://next.hanakua.mx/Notas?id=](https://next.hanakua.mx/Notas?id=)...)
    urls_hanakua_notas = [u for u in urls_sin_testigo if "hanakua.mx/Notas" in u]
    if urls_hanakua_notas:
        return urls_hanakua_notas[0]
        
    # Respaldo
    if urls_sin_testigo:
        return urls_sin_testigo[0]
        
    return ""

def crear_doc_desde_hoja(df_hoja, nombre_hoja, es_redes_sociales):
    if df_hoja is None or df_hoja.empty:
        return None

    # Auto-reparar posibles desfasamientos humanos de columnas en Excel
    df_hoja = reparar_desfase_columnas_excel(df_hoja)

    # 1. Eliminar filas con leyendas tipo 'sin notas'
    mask_sin_notas = df_hoja.apply(lambda r: any(k in str(v).lower() for v in r.values for k in ['sin notas', 'si notas', 'sin nota', 'sin registro']), axis=1)
    df_hoja = df_hoja[~mask_sin_notas].copy()
    if len(df_hoja) == 0:
        return None

    # 2. Limpieza universal de cuentas oficiales y personales del actor político
    df_filtrado, total_descartadas = limpiar_dataframe_redes_automatico(df_hoja, nombre_hoja)

    if len(df_filtrado) == 0:
        return None

    serie_fechas_raw = obtener_columna_serie(
        df_filtrado, ["Publish date", "Fecha", "Date", "Fecha de publicación", "Fecha de publicacion"]
    )
    df_filtrado["fecha_dt"] = serie_fechas_raw.apply(parsear_fecha_perfecta)

    df_filtrado = (
        df_filtrado
        .dropna(subset=["fecha_dt"])
        .sort_values(by="fecha_dt", ascending=True)
    )

    if len(df_filtrado) == 0:
        return None

    # Descartar duplicados si se fusionaron múltiples archivos
    subset_dup = [c for c in ["ID Nota", "Titulo", "Link de Nota", "Link URL Medio"] if c in df_filtrado.columns]
    if len(subset_dup) > 0:
        df_filtrado = df_filtrado.drop_duplicates(subset=subset_dup)

    df_filtrado["fecha_str"] = df_filtrado["fecha_dt"].dt.strftime("%d.%m.%2y")

    fechas_validas = df_filtrado["fecha_dt"]
    min_d = fechas_validas.min()
    max_d = fechas_validas.max()

    if min_d.strftime("%d.%m") == max_d.strftime("%d.%m"):
        periodo_texto = f"{min_d.strftime('%d')} de {MESES_ES[max_d.month]} de {max_d.year}"
    else:
        periodo_texto = f"{min_d.strftime('%d')} al {max_d.strftime('%d')} de {MESES_ES[max_d.month]} de {max_d.year}"

    df_filtrado["sentimiento_final"] = determinar_sentimiento_df(df_filtrado, nombre_hoja, es_tradicionales=not es_redes_sociales)

    positivas_cnt = len(df_filtrado[df_filtrado["sentimiento_final"].isin(["POSITIVA", "NEUTRA"])])
    negativas_cnt = len(df_filtrado[df_filtrado["sentimiento_final"] == "NEGATIVA"])
    total_cnt = len(df_filtrado)

    # Identificación estandarizada de medio
    serie_media_raw = obtener_columna_serie(
        df_filtrado,
        ["Tipo de Medio", "Fuente", "Media type", "Media Type", "Medio", "Nombre del Medio", "Canal", "Tipo de Nota"]
    )
    df_filtrado["categoria_medio_std"] = serie_media_raw.apply(estandarizar_categoria_medio) if not es_redes_sociales else "REDES SOCIALES"

    tv_cnt = (df_filtrado["categoria_medio_std"] == "TELEVISIÓN").sum() if not es_redes_sociales else 0
    rad_cnt = (df_filtrado["categoria_medio_std"] == "RADIO").sum() if not es_redes_sociales else 0
    portales_cnt = (df_filtrado["categoria_medio_std"] == "PORTALES DIGITALES").sum() if not es_redes_sociales else 0
    prensa_cnt = (df_filtrado["categoria_medio_std"] == "PRENSA LOCAL").sum() if not es_redes_sociales else 0
    columnas_cnt = (df_filtrado["categoria_medio_std"] == "COLUMNAS").sum() if not es_redes_sociales else 0

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    doc.styles["Normal"].font.name = "Verdana"
    doc.styles["Normal"].font.size = Pt(10)

    def add_run_verdana(p, text, bold=False, italic=False, size_pt=10, color_rgb=None, underline=False):
        run = p.add_run(text)
        run.font.name = "Verdana"
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size_pt)
        run.underline = underline
        if color_rgb:
            run.font.color.rgb = color_rgb
        return run

    def fondo_celda(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

    # 1. Encabezado
    p_title = doc.add_paragraph()
    add_run_verdana(p_title, nombre_hoja.upper(), bold=True, size_pt=12, color_rgb=RGBColor(0, 51, 102))

    p_per = doc.add_paragraph()
    add_run_verdana(p_per, f"PERIODO DE MEDICIÓN: {periodo_texto}", bold=True, size_pt=10)

    p_can = doc.add_paragraph()
    p_can.paragraph_format.space_after = Pt(8)
    add_run_verdana(p_can, "CANALES: PRENSA, TV, RADIO, PORTALES, REDES SOCIALES Y COLUMNAS.", bold=True, size_pt=9.5)

    # 2. Balance de Impactos
    p_bal = doc.add_paragraph()
    add_run_verdana(p_bal, "BALANCE DE IMPACTOS", bold=True, size_pt=10.5)

    t_imp = doc.add_table(rows=2, cols=3)
    t_imp.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["POSITIVA / INFORMATIVA", "NEGATIVA", "TOTAL DE IMPACTOS"]
    fills = ["E2EFDA", "FCE4D6", "D9E1F2"]

    for col_idx, (h_text, fill_color) in enumerate(zip(headers, fills)):
        cell = t_imp.cell(0, col_idx)
        fondo_celda(cell, fill_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_verdana(p, h_text, bold=True, size_pt=9.5)

    val_counts = [str(positivas_cnt), str(negativas_cnt), str(total_cnt)]
    for col_idx, val_text in enumerate(val_counts):
        cell = t_imp.cell(1, col_idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_verdana(p, val_text, bold=True, size_pt=11)

    p_tot = doc.add_paragraph()
    p_tot.paragraph_format.space_before = Pt(10)
    p_tot.paragraph_format.space_after = Pt(10)

    if es_redes_sociales:
        texto_totales = f"TOTAL NOTAS INFORMATIVAS: {total_cnt}\nREDES SOCIALES: {total_cnt}\nPORTALES DIGITALES: 0\nPRENSA LOCAL: 0\nCOLUMNAS: 0"
    else:
        partes = [f"TOTAL NOTAS INFORMATIVAS: {total_cnt}"]
        if tv_cnt > 0: partes.append(f"TELEVISIÓN: {tv_cnt}")
        if rad_cnt > 0: partes.append(f"RADIO: {rad_cnt}")
        if portales_cnt > 0: partes.append(f"PORTALES DIGITALES: {portales_cnt}")
        if prensa_cnt > 0: partes.append(f"PRENSA LOCAL: {prensa_cnt}")
        if columnas_cnt > 0: partes.append(f"COLUMNAS: {columnas_cnt}")
        if len(partes) == 1:
            partes.extend(["PORTALES DIGITALES: 0", "TELEVISIÓN: 0", "RADIO: 0", "PRENSA LOCAL: 0", "COLUMNAS: 0"])
        texto_totales = "\n".join(partes)

    add_run_verdana(p_tot, texto_totales, bold=True, size_pt=10)

    # 3. Resumen
    p_res = doc.add_paragraph()
    p_res.paragraph_format.space_before = Pt(10)
    add_run_verdana(p_res, "RESUMEN", bold=True, size_pt=11)

    temas_texto = extraer_resumen_temas_real(df_filtrado, nombre_hoja)
    for linea in temas_texto.split("\n"):
        linea_clean = linea.strip()
        if linea_clean:
            p_t = doc.add_paragraph()
            p_t.paragraph_format.space_before = Pt(1)
            p_t.paragraph_format.space_after = Pt(2)
            if "TEMAS RELEVANTES INFORMATIVOS" in linea_clean.upper():
                add_run_verdana(p_t, "Temas relevantes informativos", bold=True, size_pt=10)
            elif "TEMAS NEGATIVOS" in linea_clean.upper():
                p_t.paragraph_format.space_before = Pt(4)
                add_run_verdana(p_t, "Temas negativos", bold=True, size_pt=10, color_rgb=RGBColor(180, 0, 0))
            else:
                add_run_verdana(p_t, linea_clean, size_pt=9.5)

    # 4. Desglose
    p_des = doc.add_paragraph()
    p_des.paragraph_format.space_before = Pt(12)
    add_run_verdana(p_des, "DESGLOSE", bold=True, size_pt=11)

    orden_medios_oficial = ["TELEVISIÓN", "RADIO", "PORTALES DIGITALES", "PRENSA LOCAL", "COLUMNAS", "REDES SOCIALES"]

    for fecha_dt_val, sub_df in df_filtrado.groupby("fecha_dt", sort=True):
        fecha_item = fecha_dt_val.strftime("%d.%m.%2y")

        p_f = doc.add_paragraph()
        p_f.paragraph_format.space_before = Pt(10)
        p_f.paragraph_format.space_after = Pt(2)
        add_run_verdana(p_f, fecha_item, bold=True, size_pt=10.5, color_rgb=RGBColor(0, 51, 102))

        pos_df = sub_df[sub_df["sentimiento_final"].isin(["POSITIVA", "NEUTRA"])]
        neg_df = sub_df[sub_df["sentimiento_final"] == "NEGATIVA"]

        if es_redes_sociales or (len(pos_df) > 0 and pos_df["categoria_medio_std"].iloc[0] == "REDES SOCIALES"):
            if len(pos_df) > 0:
                p_m = doc.add_paragraph()
                p_m.paragraph_format.space_before = Pt(4)
                p_m.paragraph_format.space_after = Pt(4)
                add_run_verdana(p_m, f"REDES SOCIALES: {len(pos_df)}", bold=True, size_pt=10)

                for _, row in pos_df.iterrows():
                    autor = obtener_campo(row, ["Autor", "Author name", "Fuente", "Media name", "Programa"])
                    handle = obtener_campo(row, ["Author handle (@username)", "Handle", "Username"])
                    detalle = obtener_campo(row, ["Contenido", "Detail", "Summary", "Síntesis", "Sintesis", "Titulo", "Título", "Title", "Encabezado"])
                    link = obtener_link_inteligente(row)

                    p_a = doc.add_paragraph()
                    p_a.paragraph_format.space_before = Pt(4)
                    p_a.paragraph_format.space_after = Pt(1)
                    if handle and not handle.startswith("@"): handle = f"@{handle}"
                    add_run_verdana(p_a, f"{autor} {handle}".strip() if handle else autor, bold=True, size_pt=10)

                    p_d = doc.add_paragraph()
                    p_d.paragraph_format.space_after = Pt(2)
                    add_run_verdana(p_d, limpiar_texto(detalle), bold=False, size_pt=9.5)

                    if link:
                        p_l = doc.add_paragraph()
                        p_l.paragraph_format.space_after = Pt(6)
                        add_run_verdana(p_l, link, bold=False, size_pt=9, color_rgb=RGBColor(0, 102, 204), underline=True)

            if len(neg_df) > 0:
                p_neg_hdr = doc.add_paragraph()
                p_neg_hdr.paragraph_format.space_before = Pt(6)
                p_neg_hdr.paragraph_format.space_after = Pt(4)
                add_run_verdana(p_neg_hdr, f"NEGATIVAS: {len(neg_df)}", bold=True, size_pt=10, color_rgb=RGBColor(180, 0, 0))

                for _, row in neg_df.iterrows():
                    autor = obtener_campo(row, ["Autor", "Author name", "Fuente", "Media name", "Programa"])
                    handle = obtener_campo(row, ["Author handle (@username)", "Handle", "Username"])
                    detalle = obtener_campo(row, ["Contenido", "Detail", "Summary", "Síntesis", "Sintesis", "Titulo", "Título", "Title", "Encabezado"])
                    link = obtener_link_inteligente(row)

                    p_a = doc.add_paragraph()
                    p_a.paragraph_format.space_before = Pt(4)
                    p_a.paragraph_format.space_after = Pt(1)
                    if handle and not handle.startswith("@"): handle = f"@{handle}"
                    add_run_verdana(p_a, f"{autor} {handle}".strip() if handle else autor, bold=True, size_pt=10)

                    p_d = doc.add_paragraph()
                    p_d.paragraph_format.space_after = Pt(2)
                    add_run_verdana(p_d, limpiar_texto(detalle), bold=False, size_pt=9.5)

                    if link:
                        p_l = doc.add_paragraph()
                        p_l.paragraph_format.space_after = Pt(6)
                        add_run_verdana(p_l, link, bold=False, size_pt=9, color_rgb=RGBColor(0, 102, 204), underline=True)

        else:
            # MEDIOS TRADICIONALES
            if len(pos_df) > 0:
                for cat_nombre in orden_medios_oficial:
                    sub_pos_cat = pos_df[pos_df["categoria_medio_std"] == cat_nombre]
                    if len(sub_pos_cat) > 0:
                        p_m = doc.add_paragraph()
                        p_m.paragraph_format.space_before = Pt(4)
                        p_m.paragraph_format.space_after = Pt(4)
                        add_run_verdana(p_m, f"{cat_nombre}: {len(sub_pos_cat)}", bold=True, size_pt=10)

                        for _, row in sub_pos_cat.iterrows():
                            medio = obtener_campo(row, ["Nombre del Medio", "Fuente", "Media name", "Medio"])
                            autor = obtener_campo(row, ["Autor", "Author name", "Programa", "Conductor"])
                            hora = obtener_campo(row, ["Hora", "Hour", "Time", "Hora de Transmisión", "Hora de Transmision"])
                            titulo = obtener_campo(row, ["Titulo", "Título", "Contenido", "Detail", "Summary", "Síntesis", "Sintesis", "Encabezado", "Nota"])
                            link = obtener_link_inteligente(row)

                            p_a = doc.add_paragraph()
                            p_a.paragraph_format.space_before = Pt(4)
                            p_a.paragraph_format.space_after = Pt(1)
                            cabecera = f"{medio} - {autor}" if (autor and autor not in ["Redacción", "Staff", "Online", medio]) else medio
                            add_run_verdana(p_a, cabecera, bold=True, size_pt=10)

                            cuerpo_texto = limpiar_texto(titulo)
                            if hora and not cuerpo_texto.lower().startswith(hora.lower()):
                                cuerpo_texto = f"{hora} {cuerpo_texto}".strip()

                            p_d = doc.add_paragraph()
                            p_d.paragraph_format.space_after = Pt(2)
                            add_run_verdana(p_d, cuerpo_texto, bold=False, size_pt=9.5)

                            if link:
                                p_l = doc.add_paragraph()
                                p_l.paragraph_format.space_after = Pt(6)
                                add_run_verdana(p_l, link, bold=False, size_pt=9, color_rgb=RGBColor(0, 102, 204), underline=True)

            if len(neg_df) > 0:
                p_neg_hdr = doc.add_paragraph()
                p_neg_hdr.paragraph_format.space_before = Pt(6)
                p_neg_hdr.paragraph_format.space_after = Pt(4)
                add_run_verdana(p_neg_hdr, f"NEGATIVAS: {len(neg_df)}", bold=True, size_pt=10, color_rgb=RGBColor(180, 0, 0))

                for cat_nombre in orden_medios_oficial:
                    sub_neg_cat = neg_df[neg_df["categoria_medio_std"] == cat_nombre]
                    if len(sub_neg_cat) > 0:
                        p_sub_neg = doc.add_paragraph()
                        p_sub_neg.paragraph_format.space_before = Pt(4)
                        p_sub_neg.paragraph_format.space_after = Pt(4)
                        add_run_verdana(p_sub_neg, f"{cat_nombre}: {len(sub_neg_cat)}", bold=True, size_pt=10, color_rgb=RGBColor(180, 0, 0))

                        for _, row in sub_neg_cat.iterrows():
                            medio = obtener_campo(row, ["Nombre del Medio", "Fuente", "Media name", "Medio"])
                            autor = obtener_campo(row, ["Autor", "Author name", "Programa", "Conductor"])
                            hora = obtener_campo(row, ["Hora", "Hour", "Time", "Hora de Transmisión", "Hora de Transmision"])
                            titulo = obtener_campo(row, ["Titulo", "Título", "Contenido", "Detail", "Summary", "Síntesis", "Sintesis", "Encabezado", "Nota"])
                            link = obtener_link_inteligente(row)

                            p_a = doc.add_paragraph()
                            p_a.paragraph_format.space_before = Pt(4)
                            p_a.paragraph_format.space_after = Pt(1)
                            cabecera = f"{medio} - {autor}" if (autor and autor not in ["Redacción", "Staff", "Online", medio]) else medio
                            add_run_verdana(p_a, cabecera, bold=True, size_pt=10)

                            cuerpo_texto = limpiar_texto(titulo)
                            if hora and not cuerpo_texto.lower().startswith(hora.lower()):
                                cuerpo_texto = f"{hora} {cuerpo_texto}".strip()

                            p_d = doc.add_paragraph()
                            p_d.paragraph_format.space_after = Pt(2)
                            add_run_verdana(p_d, cuerpo_texto, bold=False, size_pt=9.5)

                            if link:
                                p_l = doc.add_paragraph()
                                p_l.paragraph_format.space_after = Pt(6)
                                add_run_verdana(p_l, link, bold=False, size_pt=9, color_rgb=RGBColor(0, 102, 204), underline=True)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- INTERFAZ STREAMLIT ---

tipo_analisis = st.radio(
    "¿Qué tipo de archivo vas a analizar?",
    [
        "Redes Sociales",
        "Medios Tradicionales / Portales / TV y Radio (Multi-Archivo)"
    ],
    index=0
)

if tipo_analisis == "Redes Sociales":
    uploaded_file = st.file_uploader(
        "Sube tu archivo Excel o CSV de Redes Sociales",
        type=["xlsx", "xls", "csv"]
    )
    actor_nombre_in = st.text_input(
        "Nombre y Partido del Actor Político",
        placeholder="ej. ALEJANDRO ARMENTA (MORENA), EDUARDO RIVERA (PAN), etc.",
    ).strip().upper()

    if uploaded_file and actor_nombre_in:
        if st.button("Generar Reporte Oficial", type="primary"):
            with st.spinner("Limpiando cuentas oficiales y evaluando sentimiento con IA desde la perspectiva del actor..."):
                try:
                    dict_h = cargar_archivo_seguro(uploaded_file)
                    df_redes = list(dict_h.values())[0]
                    buf = crear_doc_desde_hoja(df_redes, actor_nombre_in, es_redes_sociales=True)
                    if buf is not None:
                        st.success(f"¡Reporte generado exitosamente para '{actor_nombre_in}'!")
                        st.download_button(
                            label=f"📥 Descargar Reporte Word de {actor_nombre_in}",
                            data=buf,
                            file_name=f"Reporte_{actor_nombre_in.replace(' ', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.warning("El archivo no contiene notas válidas tras la limpieza automática.")
                except Exception as e:
                    st.error(f"Error procesando el archivo: {str(e)}")

else:
    uploaded_files = st.file_uploader(
        "Sube uno o varios archivos Excel (ej. Archivo de TV/Radio y Archivo de Portales Web)",
        type=["xlsx", "xls", "csv"],
        accept_multiple_files=True
    )

    if uploaded_files and len(uploaded_files) > 0:
        candidatos_unificados = {}

        for file_item in uploaded_files:
            try:
                dict_hojas = cargar_archivo_seguro(file_item)
                for h_name, df_h in dict_hojas.items():
                    if df_h is None or df_h.empty:
                        continue
                    
                    # Identificar nombre de candidato
                    if 'Menu' in df_h.columns and len(df_h['Menu'].dropna()) > 0:
                        nombre_raw = str(df_h['Menu'].dropna().iloc[0]).strip()
                    else:
                        nombre_raw = h_name
                    
                    candidato_canon = normalizar_nombre_candidato(nombre_raw)
                    
                    if candidato_canon not in candidatos_unificados:
                        candidatos_unificados[candidato_canon] = []
                    candidatos_unificados[candidato_canon].append(df_h)
            except Exception as e:
                st.warning(f"No se pudo procesar una de las hojas de {file_item.name}: {str(e)}")

        candidatos_disponibles = sorted([c for c in candidatos_unificados.keys() if c and "SIN NOTAS" not in c])

        if len(candidatos_disponibles) > 0:
            st.success(f"✅ Se cargaron exitosamente {len(uploaded_files)} archivo(s) y se detectaron {len(candidatos_disponibles)} candidatos.")
            st.markdown(f"**Candidatos detectados:** {', '.join(candidatos_disponibles)}")

            st.write("---")
            st.subheader("Opciones de Descarga:")

            col1, col2 = st.columns(2)

            with col1:
                st.markdown("### 📦 Descarga Masiva")
                if st.button("Generar y Descargar TODOS los Reportes en .ZIP", type="primary", use_container_width=True):
                    with st.spinner("Generando reportes consolidados para todos los candidatos..."):
                        zip_buffer = io.BytesIO()
                        cnt_generados = 0
                        
                        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                            for cand_name in candidatos_disponibles:
                                lista_dfs = candidatos_unificados[cand_name]
                                df_total_candidato = pd.concat(lista_dfs, ignore_index=True)

                                buf = crear_doc_desde_hoja(df_total_candidato, cand_name, es_redes_sociales=False)
                                if buf is not None:
                                    doc_bytes = buf.getvalue()
                                    fname = f"Reporte_{cand_name.replace(' ', '_')}.docx"
                                    zip_file.writestr(fname, doc_bytes)
                                    cnt_generados += 1

                        zip_buffer.seek(0)
                        if cnt_generados > 0:
                            st.success(f"¡Se generaron con éxito {cnt_generados} reportes consolidados!")
                            st.download_button(
                                label="📥 Descargar Archivo .ZIP",
                                data=zip_buffer,
                                file_name="Reportes_Monitoreo_Consolidados_Completos.zip",
                                mime="application/zip",
                                use_container_width=True
                            )
                        else:
                            st.warning("No se encontraron candidatos con notas activas.")

            with col2:
                st.markdown("### 📄 Descarga Individual")
                cand_sel = st.selectbox("Selecciona un candidato:", candidatos_disponibles)
                if st.button(f"Generar Reporte de {cand_sel}", use_container_width=True):
                    with st.spinner(f"Consolidando notas para {cand_sel}..."):
                        lista_dfs = candidatos_unificados[cand_sel]
                        df_total_candidato = pd.concat(lista_dfs, ignore_index=True)
                        
                        buf = crear_doc_desde_hoja(df_total_candidato, cand_sel, es_redes_sociales=False)
                        if buf is not None:
                            st.success(f"¡Reporte consolidado listo para '{cand_sel}'!")
                            st.download_button(
                                label=f"📥 Descargar Word de {cand_sel}",
                                data=buf,
                                file_name=f"Reporte_{cand_sel.replace(' ', '_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        else:
                            st.warning(f"El candidato '{cand_sel}' no contiene notas registradas.")
        else:
            st.warning("No se detectaron candidatos con notas válidas en los archivos seleccionados.")
